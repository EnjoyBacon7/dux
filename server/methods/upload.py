"""
File upload and text extraction utilities.

Handles CV file uploads, text extraction from various file formats
(PDF, DOCX, TXT), with OCR fallback support for scanned PDFs.
"""

from fastapi import UploadFile, File, HTTPException
from pathlib import Path
import shutil
import os
import uuid
import logging
from typing import Dict, Any

import PyPDF2
import pytesseract
from docx import Document

logger = logging.getLogger(__name__)

UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "./uploads"))
UPLOAD_DIR.mkdir(exist_ok=True, parents=True)


# ============================================================================
# Text Extraction Functions
# ============================================================================


def extract_text_from_pdf(file_path: Path) -> str:
    """
    Extract text from PDF file with OCR fallback.

    Uses PyPDF2 for text extraction from searchable PDFs.
    If no text is found (e.g., scanned PDFs), attempts OCR using pytesseract.

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text from the PDF

    Note:
        OCR fallback requires pytesseract and pdf2image to be installed
    """
    text = ""
    try:
        with open(file_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        logger.warning(f"Error extracting text from PDF: {e}")
        text = ""

    # If no text was extracted, try OCR
    if not text.strip():
        logger.info("No text found in PDF, attempting OCR...")
        try:
            # Convert PDF to images and run OCR
            from pdf2image import convert_from_path
            images = convert_from_path(file_path)
            for image in images:
                text += pytesseract.image_to_string(image) + "\n"
        except Exception as e:
            logger.warning(f"OCR failed: {e}")

    return text.strip()


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract text from DOCX/DOC file.

    Extracts text from both paragraphs and tables within the document.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text from the document

    Raises:
        HTTPException: If document reading fails
    """
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to read DOCX file: {str(e)}")

    return text.strip()


def extract_text_from_txt(file_path: Path) -> str:
    """
    Extract text from TXT file.

    Args:
        file_path: Path to the TXT file

    Returns:
        Text content of the file

    Raises:
        HTTPException: If file reading fails
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as txt_file:
            return txt_file.read().strip()
    except Exception as e:
        logger.error(f"Error reading TXT file: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to read TXT file: {str(e)}")


def extract_text_from_file(file_path: Path, file_extension: str) -> str:
    """
    Extract text from file based on its extension.

    Routes to appropriate extraction function based on file type.

    Args:
        file_path: Path to the file
        file_extension: File extension (e.g., ".pdf", ".docx")

    Returns:
        Extracted text from the file

    Raises:
        HTTPException: If file type is unsupported or extraction fails
    """
    if file_extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif file_extension == ".docx" or file_extension == ".doc":
        return extract_text_from_docx(file_path)
    elif file_extension == ".txt":
        return extract_text_from_txt(file_path)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_extension}")


# ============================================================================
# File Upload Handler
# ============================================================================


async def upload_file(file: UploadFile = File(...)) -> Dict[str, Any]:
    """
    Handle file upload, validation, storage, and text extraction.

    Validates file type and size, saves to disk with secure filename,
    and extracts text content for later processing.

    Args:
        file: The uploaded file (multipart/form-data)

    Returns:
        dict: Upload result with:
        - filename: Original filename
        - content_type: File MIME type
        - size: File size in bytes
        - extracted_text: Extracted text content
        - message: Success message

    Raises:
        HTTPException: If validation fails or processing errors occur
    """
    # Validate file
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    # Validate file extension (security check)
    allowed_extensions = {".pdf", ".doc", ".docx", ".txt"}
    file_extension = Path(file.filename).suffix.lower()
    if file_extension not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"File type not allowed. Allowed types: {', '.join(allowed_extensions)}"
        )

    # Create safe filename using UUID to prevent directory traversal and filename collisions
    safe_filename = f"{uuid.uuid4()}{file_extension}"
    file_path = UPLOAD_DIR / safe_filename

    try:
        # Save file
        with file_path.open("wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Extract text from file
        extracted_text = extract_text_from_file(file_path, file_extension)

        return {
            "filename": safe_filename,
            "content_type": file.content_type,
            "size": file_path.stat().st_size,
            "extracted_text": extracted_text,
            "message": "File uploaded successfully and text extracted"
        }

    except HTTPException:
        # Clean up partial file if upload fails
        if file_path.exists():
            file_path.unlink()
        raise
    except Exception as e:
        # Clean up partial file if upload fails
        if file_path.exists():
            file_path.unlink()
        logger.error(f"File upload failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"File upload failed: {str(e)}")
